'''职准星 - 报告生成模块

接收 session 完整测评、岗位推荐、差距分析全量数据，生成：
  1. 标准结构化 Markdown 报告（7 章节，含评分/红黄绿灯/优化对比/话术包）
  2. 基于 Markdown 渲染 docx 二进制流
  3. 对话内展示摘要（build_report_summary）

内置 REPORT_CACHE_MINUTES 缓存，相同简历+JD 直接复用。
统一入口函数 generate_full_report(session)。
'''

import os
import re
import json
import hashlib
import logging
import io
from typing import Any, Dict, Optional, Tuple
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..core import config
from ..core.llm import llm_client
from ..core.prompts import REPORT_HEADER_TEMPLATE, REPORT_SUMMARY_TEMPLATE
from ..modules.gap_analysis import run_gap_analysis
from ..models.session import Session

logger = logging.getLogger(__name__)


# ============================================================
# 报告缓存 {cache_key: (timestamp, markdown, docx_bytes, filename)}
# ============================================================

_report_cache: Dict[str, Tuple[datetime, str, bytes, str]] = {}


def _build_cache_key(session: Session) -> str:
    '''基于 session 中简历+JD 的 hash 生成缓存 key。'''
    raw = (session.final_resume or '')[:200] + (session.final_jd or '')[:200]
    return hashlib.md5(raw.encode()).hexdigest()


# ============================================================
# 统一入口函数
# ============================================================

async def generate_full_report(session: Session) -> dict:
    '''
    报告生成统一入口函数。

    流程：
      1. 检查缓存（10 分钟内同 JD+简历复用）
      2. 从 session 读取 gap_analysis_result / optimized_resume_result / counterattack_result
      3. 若数据为空，先执行差距分析工作流
      4. 拼接 Markdown 报告
      5. 渲染 docx 二进制流
      6. 写入缓存
      7. 返回 { markdown, docx_bytes, filename, summary, report_url }

    Args:
        session: 当前会话（含 final_resume, final_jd, gap_analysis_result 等）

    Returns:
        {
            'markdown':   str,    # 完整 Markdown 报告
            'docx_bytes': bytes,  # docx 文件二进制流
            'filename':   str,    # 文件名
            'summary':    str,    # 对话内展示摘要
            'report_url': str,    # 下载链接路径
            'cached':     bool,   # 是否命中缓存
        }
    '''
    # ---- 1. 缓存检查 ----
    cache_key = _build_cache_key(session)
    now = datetime.now()

    if cache_key in _report_cache:
        cached_time, md, docx_bytes, fname = _report_cache[cache_key]
        if now - cached_time < timedelta(minutes=config.REPORT_CACHE_MINUTES):
            logger.info(f'会话 {session.id}: 命中报告缓存 ({fname})')
            summary = build_report_summary(session.gap_analysis_result or {})
            return {
                'markdown': md,
                'docx_bytes': docx_bytes,
                'filename': fname,
                'summary': summary,
                'report_url': f'/api/report/{fname}',
                'cached': True,
            }

    # ---- 2. 确保差距分析结果存在 ----
    gap_data = session.gap_analysis_result or {}
    opt_data = session.optimized_resume_result or {}
    cta_data = session.counterattack_result or {}

    if not gap_data or not opt_data or not cta_data:
        logger.info(f'会话 {session.id}: 差距分析数据为空，自动执行差距分析')
        results = await run_gap_analysis(
            session,
            session.final_resume or '',
            session.final_jd or '',
        )
        gap_data = results.get('gap_data', {})
        opt_data = results.get('optimized_data', {})
        cta_data = results.get('counterattack_data', {})
        session.gap_analysis_result = gap_data
        session.optimized_resume_result = opt_data
        session.counterattack_result = cta_data

    # ---- 3. 拼接 Markdown ----
    user_name = _extract_user_name(session.final_resume or '')
    markdown = build_report_markdown(gap_data, opt_data, cta_data, user_name)

    # ---- 4. 渲染 docx 二进制流 ----
    docx_bytes = render_docx_bytes(markdown, user_name)

    # ---- 5. 生成文件名 ----
    timestamp = now.strftime('%Y%m%d_%H%M%S')
    name_part = f'_{user_name}' if user_name else ''
    safe_name = re.sub(r'[\\/:*?"<>|]', '_', name_part)
    filename = f'\u804c\u51c6\u661f\u5dee\u8ddd\u5206\u6790\u62a5\u544a{safe_name}_{timestamp}.docx'

    # 同时写入磁盘（供 FileResponse 下载）
    os.makedirs(config.REPORT_DIR, exist_ok=True)
    filepath = os.path.join(config.REPORT_DIR, filename)
    with open(filepath, 'wb') as f:
        f.write(docx_bytes)

    session.report_filename = filename
    session.report_generated_at = now

    # ---- 6. 写入缓存 ----
    _report_cache[cache_key] = (now, markdown, docx_bytes, filename)

    # ---- 7. 组装返回 ----
    summary = build_report_summary(gap_data)
    report_url = f'/api/report/{filename}'

    logger.info(
        f'会话 {session.id}: 报告生成完成 | 文件名={filename} | '
        f'docx={len(docx_bytes)}bytes'
    )

    return {
        'markdown': markdown,
        'docx_bytes': docx_bytes,
        'filename': filename,
        'summary': summary,
        'report_url': report_url,
        'cached': False,
    }


# ============================================================
# Markdown 报告拼接
# ============================================================

def build_report_markdown(
    gap_data: dict,
    opt_data: dict,
    cta_data: dict,
    user_name: str = '',
) -> str:
    '''
    将 3 个 JSON 对象拼接为完整 Markdown 报告（7 章节）。

    Args:
        gap_data: 差距诊断结果（match_score + gap_analysis + top3_fatal_gaps）
        opt_data: 简历优化结果（optimized_resume + changes + highlights_added）
        cta_data: 逆袭话术结果（counterattack_scripts + predicted_questions + tips）
        user_name: 用户姓名（可选，写入页眉）

    Returns:
        Markdown 格式报告文本（所有字段有 .get() + or [] 容错）
    '''
    # ---- 评分 ----
    match_score = gap_data.get('match_score', {}) or {}
    total = match_score.get('total', 'N/A')
    edu = match_score.get('education_score', 'N/A')
    exp = match_score.get('experience_score', 'N/A')
    skill = match_score.get('skill_score', 'N/A')
    calc_detail = match_score.get('calculation_detail', '')

    # ---- 红黄绿灯 ----
    gap_analysis = gap_data.get('gap_analysis', {}) or {}
    green_items = gap_analysis.get('green_items', []) or []
    yellow_items = gap_analysis.get('yellow_items', []) or []
    red_items = gap_analysis.get('red_items', []) or []
    top3 = gap_data.get('top3_fatal_gaps', []) or []
    priority_fix = gap_data.get('priority_fix_order', []) or []

    # ---- 优化 ----
    optimized_resume = opt_data.get('optimized_resume', '（无优化内容）') or '（无优化内容）'
    changes = opt_data.get('changes', []) or []
    highlights_added = opt_data.get('highlights_added', []) or []

    # ---- 拼装 ----
    md: list[str] = []

    name_section = f' - {user_name}' if user_name else ''
    md.append(f'# 职准星 · 简历-JD 差距分析报告{name_section}\n')

    # 1. 匹配评分
    md.append('## \U0001f4ca 1. 匹配评分')
    md.append(f'- **总分**：{total} / 100')
    md.append(f'- \U0001f393 教育匹配：{edu} 分')
    md.append(f'- \U0001f4bc 经验匹配：{exp} 分')
    md.append(f'- \U0001f527 技能匹配：{skill} 分')
    if calc_detail:
        md.append(f'- **计算明细**：{calc_detail}')
    md.append('')

    # 2. 差距诊断
    md.append('## \U0001f6a6 2. 差距诊断（红黄绿灯）')
    if green_items:
        md.append('### \U0001f7e2 达标项')
        for i in green_items:
            md.append(f'- **{i.get("item", "")}**：{i.get("detail", "")}')
            if i.get('evidence'):
                md.append(f'  - 证据：{i.get("evidence", "")}')
    if yellow_items:
        md.append('\n### \U0001f7e1 需优化项')
        for i in yellow_items:
            md.append(f'- **{i.get("item", "")}**：{i.get("detail", "")}')
            if i.get('suggestion'):
                md.append(f'  - 建议：{i.get("suggestion", "")}')
    if red_items:
        md.append('\n### \U0001f534 严重缺失项')
        for i in red_items:
            md.append(
                f'- **{i.get("item", "")}**：{i.get("detail", "")}'
                f'（紧急度：{i.get("urgency", "")}）'
            )
    md.append('')

    # 3. Top3 致命差距
    md.append('## \u26a0\ufe0f 3. Top3 致命差距 + 面试官追问')
    if top3:
        for idx, i in enumerate(top3, 1):
            md.append(f'### {idx}. {i.get("gap", "")}')
            md.append(f'- **致命原因**：{i.get("why_fatal", "")}')
            md.append(f'- **面试可能问**：{i.get("interview_question", "")}')
            md.append(f'- **逆袭方向**：{i.get("counterattack_hint", "")}')
    md.append('')

    # 4. 简历优化
    md.append('## \U0001f4dd 4. 简历定向优化（关键修改对比）')
    if highlights_added:
        md.append('### \u2728 新增高光项')
        for h in highlights_added:
            md.append(f'- {h}')
        md.append('')
    if changes:
        md.append('### 关键修改对照表')
        md.append('| 原文 | 优化后 | 修改原因 | 优先级 |')
        md.append('|:-----|:-------|:---------|:------|')
        for c in changes:
            o = (c.get('original') or '').replace('\n', ' ').replace('|', '/')[:200]
            n = (c.get('optimized') or '').replace('\n', ' ').replace('|', '/')[:200]
            r = (c.get('reason') or '').replace('|', '/')[:120]
            p = c.get('priority', '')
            md.append(f'| {o} | {n} | {r} | {p} |')
    md.append('\n### 完整优化版简历')
    md.append('```')
    md.append(optimized_resume)
    md.append('```')
    md.append('')

    # 5. 逆袭话术（无内容时整个章节不输出）
    scripts = cta_data.get('counterattack_scripts', []) or []
    if scripts:
        md.append('## \U0001f4aa 5. 逆袭话术包')
        for s in scripts:
            gap_name = s.get('gap', '短板') or '短板'
            md.append(f'### 针对「{gap_name}」的应对')
            exp_ver = s.get('script_with_experience', '') or ''
            no_exp_ver = s.get('script_no_experience', '') or ''
            one_liner = s.get('script_one_liner', '') or ''
            if exp_ver:
                md.append(f'- **有经验版**：{exp_ver}')
            if no_exp_ver:
                md.append(f'- **没经验版**：{no_exp_ver}')
            if one_liner:
                md.append(f'- **一句话金句**：{one_liner}')
            md.append('')

    # 6. 预测面试题（无内容时整个章节不输出）
    predicted_questions = cta_data.get('predicted_interview_questions', []) or []
    if predicted_questions:
        md.append('## \U0001f52e 6. 预测面试题及答题策略')
        for q in predicted_questions:
            q_text = q.get('question', '') or ''
            why = q.get('why_predicted', '') or ''
            strategy = q.get('answer_strategy', '') or ''
            if q_text:
                md.append(f'- **Q**：{q_text}')
                if why:
                    md.append(f'  - 依据：{why}')
                if strategy:
                    md.append(f'  - 策略：{strategy}')
            md.append('')

    # 7. 面试建议（无内容时整个章节不输出）
    tips_list = cta_data.get('interview_tips', []) or []
    if tips_list:
        md.append('## \U0001f3af 7. 通用面试建议')
        for t in tips_list:
            if t:
                md.append(f'- {t}')
        md.append('')

    md.append('---\n*报告由 职准星 自动生成，祝你面试顺利！\U0001f3af*')

    return '\n'.join(md)


# ============================================================
# docx 渲染（二进制流）
# ============================================================

def render_docx_bytes(markdown_text: str, user_name: str = '') -> bytes:
    '''
    将 Markdown 报告渲染为 docx 二进制流（不写磁盘）。

    Args:
        markdown_text: Markdown 报告文本
        user_name:     用户姓名（写入页眉）

    Returns:
        docx 文件的二进制字节流
    '''
    doc = Document()

    # ---- 默认样式 ----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---- 页眉 ----
    if user_name:
        header = doc.sections[0].header
        hp = header.paragraphs[0]
        hp.text = f'\u804c\u51c6\u661f\u5dee\u8ddd\u5206\u6790\u62a5\u544a - {user_name}'
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if hp.runs:
            hp.runs[0].font.size = Pt(9)
            hp.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # ---- 逐行解析 Markdown ----
    lines = markdown_text.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 一级标题（# ）
        if stripped.startswith('# ') and not stripped.startswith('## '):
            h = doc.add_heading(stripped[2:], level=0)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0, 51, 102)

        # 二级标题（## ）
        elif stripped.startswith('## ') and not stripped.startswith('### '):
            doc.add_heading(stripped[3:], level=1)

        # 三级标题（### ）
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)

        # 分割线
        elif stripped.startswith('---'):
            doc.add_paragraph('\u2500' * 50)

        # 表格行
        elif stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if cells:
                p = doc.add_paragraph(' | '.join(cells))
                # 表格分隔行（|---|---|）跳过
                if not all(set(c) <= {'-', ':', ' '} for c in cells):
                    for run in p.runs:
                        run.font.size = Pt(10)

        # 列表项
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            for run in p.runs:
                if run.text:
                    run.font.size = Pt(11)

        # 代码块标记
        elif stripped.startswith('```'):
            continue

        # 普通段落
        else:
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
            p = doc.add_paragraph(text)

    # ---- 输出二进制 ----
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# 对话内展示摘要
# ============================================================

def build_report_summary(gap_data: dict) -> str:
    '''
    构建对话内展示的报告摘要（短文本，含评分+红黄绿灯摘要）。

    Args:
        gap_data: 差距诊断结果

    Returns:
        摘要文本（对话内直接展示）
    '''
    match_score = gap_data.get('match_score', {}) or {}
    total = match_score.get('total', 'N/A')

    gap_analysis = gap_data.get('gap_analysis', {}) or {}
    red_items = gap_analysis.get('red_items', []) or []
    yellow_items = gap_analysis.get('yellow_items', []) or []
    green_items = gap_analysis.get('green_items', []) or []

    lines = [
        '\u5df2\u5b8c\u6210\u7b80\u5386\u4e0e\u5c97\u4f4dJD\u5339\u914d\u5206\u6790\uff01',
        '',
        f'\U0001f4ca **\u5339\u914d\u8bc4\u5206\uff1a{total} / 100**',
        '',
    ]

    if red_items:
        lines.append(f'\U0001f534 \u4e25\u91cd\u7f3a\u5931\uff1a{red_items[0].get("item", "")}')
    if yellow_items:
        lines.append(f'\U0001f7e1 \u9700\u4f18\u5316\uff1a{yellow_items[0].get("item", "")}')
    if green_items:
        lines.append(f'\U0001f7e2 \u8fbe\u6807\uff1a{green_items[0].get("item", "")}')

    lines += [
        '',
        '\U0001f4aa \u9006\u88ad\u8bdd\u672f\u5df2\u751f\u6210\uff08\u6bcf\u4e2a\u5dee\u8ddd3\u7248\u672c\uff09',
        '\U0001f52e \u9762\u8bd5\u9898\u5df2\u9884\u6d4b\uff08\u542b\u7b54\u9898\u7b56\u7565\uff09',
        '\U0001f4dd \u7b80\u5386\u5df2\u4f18\u5316\uff08\u9010\u6761\u5bf9\u6bd4+\u5b8c\u6574\u7248\uff09',
        '',
        '\U0001f4c4 \u5b8c\u6574\u62a5\u544a\u5df2\u751f\u6210\uff0c\u70b9\u51fb\u4e0b\u65b9\u94fe\u63a5\u4e0b\u8f7d\uff1a',
    ]

    return '\n'.join(lines)


# ============================================================
# 工具函数
# ============================================================

def _extract_user_name(resume_text: str) -> str:
    '''
    从简历文本中简单提取姓名（用于写入报告页眉）。

    提取规则：简历前 200 字中，匹配 "姓名：xxx" 或 "姓名 xxx" 模式。
    '''
    match = re.search(r'\u59d3\u540d[\uff1a:]\\s*(\\S{2,4})', resume_text[:200])
    return match.group(1) if match else ''


def clear_cache() -> None:
    '''清空报告缓存（会话重置时调用）。'''
    _report_cache.clear()
    logger.info('报告缓存已清空')


# ============================================================
# 兼容旧接口
# ============================================================

def generate_docx_report(markdown_text: str, output_dir: str = None, user_name: str = '') -> str:
    '''
    兼容 main.py 旧接口：生成 docx 文件并返回路径。

    内部调用 render_docx_bytes() 生成二进制流后写入磁盘。
    '''
    output_dir = output_dir or config.REPORT_DIR
    os.makedirs(output_dir, exist_ok=True)

    docx_bytes = render_docx_bytes(markdown_text, user_name)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    name_suffix = f'_{user_name}' if user_name else ''
    safe_name = re.sub(r'[\\/:*?"<>|]', '_', name_suffix)
    filename = f'\u804c\u51c6\u661f\u5dee\u8ddd\u5206\u6790\u62a5\u544a{safe_name}_{timestamp}.docx'
    filepath = os.path.join(output_dir, filename)

    with open(filepath, 'wb') as f:
        f.write(docx_bytes)

    logger.info(f'报告已生成: {filepath}')
    return filepath
